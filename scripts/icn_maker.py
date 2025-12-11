"""
ICN Maker - Generate ICN labels for images in DOCX files
"""
import sys
import os
from docx import Document

def generate_icn_code(dmc_code, kpc, xyz, sq, icv, issue, sec):
    """Generate ICN code from DMC code."""
    parts = dmc_code.split("-")
    
    # Validate DMC code has enough parts
    if len(parts) < 7:
        # If not a valid DMC code, create a simple ICN format
        return f"ICN-{dmc_code}-{kpc}-{xyz}-{sq}-{icv}-{issue}-{sec}"
    
    # Standard DMC format processing
    try:
        up_to_unit = f"ICN-{parts[1]}-{parts[2]}-" + "".join(parts[3:-4]) + f"-{kpc}-{xyz}-{sq}-{icv}-{issue}-{sec}"
        return up_to_unit
    except IndexError:
        # Fallback to simple format
        return f"ICN-{dmc_code}-{kpc}-{xyz}-{sq}-{icv}-{issue}-{sec}"

def paragraph_has_image(paragraph):
    """Check if a paragraph contains an image."""
    try:
        xml = paragraph._element.xml
        return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("<wp:inline" in xml)
    except:
        return False

def generate_icn_labels(input_dir, output_dir, params):
    """Generate ICN labels for images in DOCX files."""
    os.makedirs(output_dir, exist_ok=True)
    
    current_sq = int(params['sq_start'])
    pad_len = len(params['sq_start'])
    
    files_processed = 0
    total_icns_generated = 0
    
    print(f"Starting ICN generation...", file=sys.stderr)
    print(f"Parameters: KPC={params['kpc']}, XYZ={params['xyz']}, Start={params['sq_start']}", file=sys.stderr)
    
    for filename in os.listdir(input_dir):
        if not filename.lower().endswith('.docx') or filename.startswith('~'):
            continue
        
        input_path = os.path.join(input_dir, filename)
        dmc_code = os.path.splitext(filename)[0]
        
        print(f"\nProcessing: {filename}", file=sys.stderr)
        print(f"  DMC Code: {dmc_code}", file=sys.stderr)
        
        try:
            doc = Document(input_path)
            paragraphs = doc.paragraphs[:]
            i = 0
            file_icn_count = 0
            
            while i < len(paragraphs):
                para = paragraphs[i]
                has_image = paragraph_has_image(para)
                
                if has_image:
                    icn = generate_icn_code(
                        dmc_code, params['kpc'], params['xyz'],
                        str(current_sq).zfill(pad_len), params['icv'],
                        params['issue'], params['sec']
                    )
                    current_sq += 1
                    file_icn_count += 1
                    total_icns_generated += 1
                    
                    print(f"  Generated ICN #{file_icn_count}: {icn}", file=sys.stderr)
                    
                    if icn:
                        para._p.addnext(doc.add_paragraph(icn)._p)
                        paragraphs = doc.paragraphs[:]
                        i += 1
                i += 1
            
            output_path = os.path.join(output_dir, filename)
            doc.save(output_path)
            files_processed += 1
            
            print(f"  ✓ Saved with {file_icn_count} ICN labels", file=sys.stderr)
            
        except Exception as e:
            print(f"  ✗ Failed to process {filename}: {str(e)}", file=sys.stderr)
            continue
    
    print(f"\n=== Summary ===", file=sys.stderr)
    print(f"Files processed: {files_processed}", file=sys.stderr)
    print(f"Total ICNs generated: {total_icns_generated}", file=sys.stderr)

if __name__ == "__main__":
    if len(sys.argv) != 9:
        print("Usage: icn_maker.py <input_dir> <output_dir> <kpc> <xyz> <sq_start> <icv> <issue> <sec>")
        sys.exit(1)
    
    input_dir = sys.argv[1]
    output_dir = sys.argv[2]
    params = {
        'kpc': sys.argv[3],
        'xyz': sys.argv[4],
        'sq_start': sys.argv[5],
        'icv': sys.argv[6],
        'issue': sys.argv[7],
        'sec': sys.argv[8]
    }
    
    try:
        generate_icn_labels(input_dir, output_dir, params)
        print("SUCCESS")
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        sys.exit(1)
