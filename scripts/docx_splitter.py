#!/usr/bin/env python3
"""
Standalone DOCX splitter by heading
Usage: python docx_splitter.py <input.docx> <output_dir> [heading_style]
"""

import sys
import os
import io
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

def iter_block_items(parent):
    """Yield each paragraph and table child within parent element."""
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

def split_docx_by_heading(input_path, output_dir, heading_style='Heading 1'):
    """Split DOCX file into multiple files based on heading style."""
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document(input_path)
    blocks = list(iter_block_items(doc))
    
    if not blocks:
        return 0
    
    # Find section indices
    section_indices = []
    first_block = blocks[0]
    
    if not (isinstance(first_block, Paragraph) and 
            getattr(first_block, 'style', None) and 
            getattr(first_block.style, 'name', None) == heading_style):
        section_indices.append(0)
    
    for i, block in enumerate(blocks):
        if (isinstance(block, Paragraph) and 
            getattr(block, 'style', None) and 
            getattr(block.style, 'name', None) == heading_style):
            section_indices.append(i)
    
    section_indices.append(len(blocks))
    total_sections = len(section_indices) - 1
    
    if total_sections <= 0:
        return 0
    
    # Create output files
    for i in range(total_sections):
        start, end = section_indices[i], section_indices[i + 1]
        section_blocks = blocks[start:end]
        
        if not section_blocks:
            continue
        
        # Get title from first heading
        first_block = section_blocks[0]
        if (isinstance(first_block, Paragraph) and 
            getattr(first_block, 'style', None) and 
            getattr(first_block.style, 'name', None) == heading_style):
            title = first_block.text.strip() or f"Section_{i+1}"
        else:
            title = f"Section_{i+1}"
        
        # Create safe filename
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip()[:50]
        if not safe_title:
            safe_title = f"Section_{i+1}"
        
        # Create new document
        new_doc = Document()
        
        # Copy page setup
        try:
            s1 = doc.sections[0]
            s2 = new_doc.sections[-1]
            s2.page_width = s1.page_width
            s2.page_height = s1.page_height
            s2.left_margin = s1.left_margin
            s2.right_margin = s1.right_margin
            s2.top_margin = s1.top_margin
            s2.bottom_margin = s1.bottom_margin
        except:
            pass
        
        # Copy blocks with full formatting and images
        for block in section_blocks:
            if isinstance(block, Paragraph):
                # Copy paragraph with all runs (preserves images and formatting)
                new_p = new_doc.add_paragraph()
                
                # Copy paragraph style
                try:
                    if getattr(block, 'style', None):
                        new_p.style = block.style
                except:
                    pass
                
                # Copy all runs (preserves text formatting and images)
                for run in block.runs:
                    new_run = new_p.add_run(run.text)
                    
                    # Copy run formatting
                    try:
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                    except:
                        pass
                    
                    # Copy images - need to copy actual image data
                    try:
                        # Check if run contains images
                        for inline_shape in run._element.findall('.//{*}blip'):
                            # Get the image relationship ID
                            embed_id = inline_shape.get(qn('r:embed'))
                            if embed_id:
                                # Get the image part from source document
                                image_part = doc.part.related_parts[embed_id]
                                image_bytes = image_part.blob
                                
                                # Add image to new document
                                new_run.add_picture(io.BytesIO(image_bytes))
                    except Exception as e:
                        # If image copying fails, continue without it
                        pass
                
            elif isinstance(block, Table):
                # Copy table structure
                try:
                    new_table = new_doc.add_table(rows=len(block.rows), cols=len(block.columns))
                    
                    # Copy table content
                    for i, row in enumerate(block.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                except:
                    pass
        
        # Save output file
        out_path = os.path.join(output_dir, f"{i+1:02d}_{safe_title}.docx")
        new_doc.save(out_path)
    
    return total_sections

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python docx_splitter.py <input.docx> <output_dir> [heading_style]", file=sys.stderr)
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_dir = sys.argv[2]
    heading_style = sys.argv[3] if len(sys.argv) > 3 else 'Heading 1'
    
    try:
        count = split_docx_by_heading(input_file, output_dir, heading_style)
        print(f"SUCCESS: Split into {count} files")
        sys.exit(0)
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        sys.exit(1)
